"""
generar_informe.py

Construye el informe del proyecto en formato Word (.docx) con python-docx.

Uso:
    python docs/generar_informe.py
"""

from __future__ import annotations

import os

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Pt, RGBColor, Inches

RAIZ = os.path.dirname(os.path.dirname(os.path.abspath(__file__)))
CARPETA_GRAFICAS = os.path.join(RAIZ, "graficas")
SALIDA = os.path.join(RAIZ, "docs", "Informe_Sistema_Gestion_Inventarios.docx")

GRIS = RGBColor(0x44, 0x44, 0x44)
AZUL = RGBColor(0x1F, 0x4E, 0x79)


# --------------------------------------------------------------------- #
# Utilidades de formato                                                 #
# --------------------------------------------------------------------- #
def p(doc, texto, justificar=True, italic=False, size=11):
    par = doc.add_paragraph()
    run = par.add_run(texto)
    run.italic = italic
    run.font.size = Pt(size)
    if justificar:
        par.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    return par


def codigo(doc, texto):
    par = doc.add_paragraph()
    par.paragraph_format.left_indent = Inches(0.3)
    run = par.add_run(texto)
    run.font.name = "Consolas"
    run.font.size = Pt(9)
    run.font.color.rgb = RGBColor(0x1A, 0x1A, 0x1A)
    return par


def captura(doc, descripcion):
    par = doc.add_paragraph()
    par.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = par.add_run(f"[ Espacio para captura de pantalla: {descripcion} ]")
    run.italic = True
    run.font.size = Pt(10)
    run.font.color.rgb = GRIS
    return par


def imagen(doc, nombre_archivo, ancho=6.0):
    ruta = os.path.join(CARPETA_GRAFICAS, nombre_archivo)
    if os.path.exists(ruta):
        doc.add_picture(ruta, width=Inches(ancho))
        doc.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER


def pie_figura(doc, texto):
    par = doc.add_paragraph()
    par.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = par.add_run(texto)
    run.italic = True
    run.font.size = Pt(9)
    run.font.color.rgb = GRIS


def h1(doc, texto):
    doc.add_heading(texto, level=1)


def h2(doc, texto):
    doc.add_heading(texto, level=2)


# ===================================================================== #
# Construcción del documento                                            #
# ===================================================================== #
def construir():
    doc = Document()

    # Estilo base
    estilo = doc.styles["Normal"]
    estilo.font.name = "Calibri"
    estilo.font.size = Pt(11)

    # ----------------------------- PORTADA ----------------------------- #
    for _ in range(6):
        doc.add_paragraph()
    t = doc.add_paragraph()
    t.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = t.add_run("[ ESPACIO RESERVADO PARA PORTADA PERSONALIZADA ]")
    r.italic = True
    r.font.color.rgb = GRIS

    titulo = doc.add_paragraph()
    titulo.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = titulo.add_run("Sistema de Gestión de Inventarios para una Tienda de Abarrotes")
    r.bold = True
    r.font.size = Pt(22)
    r.font.color.rgb = AZUL

    sub = doc.add_paragraph()
    sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = sub.add_run("Proyecto final — Certificación en Python (nivel intermedio)")
    r.font.size = Pt(13)
    r.italic = True

    for _ in range(8):
        doc.add_paragraph()
    aut = doc.add_paragraph()
    aut.alignment = WD_ALIGN_PARAGRAPH.CENTER
    aut.add_run("Autor: _______________________\nFecha: junio de 2026")
    doc.add_page_break()

    # --------------------------- 1. INTRODUCCIÓN ----------------------- #
    h1(doc, "1. Introducción")
    p(doc, "El presente informe documenta el diseño, desarrollo y evaluación de un "
            "sistema de gestión de inventarios orientado a tiendas de abarrotes de "
            "pequeña escala, construido íntegramente en lenguaje Python. El proyecto "
            "integra los conocimientos abordados durante la certificación: "
            "fundamentos del lenguaje, programación orientada a objetos, estructuras "
            "de datos avanzadas, manipulación y análisis de datos, así como buenas "
            "prácticas de codificación y control de versiones. La elección de Python "
            "se justifica por su legibilidad, su amplia adopción y su ecosistema de "
            "bibliotecas para el análisis de datos, características que lo posicionan "
            "como uno de los lenguajes de mayor crecimiento en la actualidad "
            "(Van Rossum & Drake, 2009).")
    p(doc, "El documento se organiza en cinco bloques: el planteamiento del problema "
            "con su justificación basada en datos oficiales; la descripción de una "
            "base de datos de ejemplo junto con su análisis exploratorio y limpieza; "
            "la evaluación de las técnicas y funciones de programación utilizadas; la "
            "complementación referente al control de datos y la persistencia; y, "
            "finalmente, las instrucciones de uso y la conclusión. Cada apartado "
            "busca demostrar tanto el dominio técnico como la reflexión crítica sobre "
            "la utilidad del sistema en un contexto real (McKinney, 2010).")

    # ------------------- 2. PLANTEAMIENTO DEL PROBLEMA ----------------- #
    h1(doc, "2. Planteamiento del problema")

    h2(doc, "2.1 Descripción de la problemática")
    p(doc, "Las micro y pequeñas empresas constituyen el tejido productivo "
            "predominante de México. De acuerdo con los Censos Económicos 2019, el "
            "99.8 % de los establecimientos del país pertenece a la categoría de "
            "micro, pequeñas y medianas empresas (mipymes), unidades que se "
            "caracterizan por una elevada movilidad en cuanto a aperturas y cierres "
            "(INEGI, 2024a). Dentro de este universo, las tiendas de abarrotes son "
            "uno de los formatos más extendidos y, a la vez, uno de los que opera con "
            "menores recursos administrativos y tecnológicos (INEGI, 2024a).")
    p(doc, "Un problema recurrente en estos negocios es la ausencia de un control "
            "sistemático del inventario. La gestión de existencias suele realizarse "
            "de memoria o mediante anotaciones manuales, lo que genera errores de "
            "captura, productos duplicados, desabastos no detectados a tiempo y "
            "decisiones de compra basadas en la intuición y no en datos. Esta "
            "carencia de herramientas de administración se ha identificado como una "
            "de las causas estructurales que limitan el crecimiento y la "
            "supervivencia de las pymes mexicanas (Ramírez Castillejo & Hernández "
            "Pérez, 2023).")
    p(doc, "La situación adquiere mayor relevancia al considerar la alta mortalidad "
            "empresarial del país. El Estudio sobre la Demografía de los Negocios "
            "2023 estimó que, entre mayo de 2019 y mayo de 2023, nacieron 1.7 "
            "millones de establecimientos mipymes y murieron 1.4 millones; expresado "
            "como tasa mensual, por cada 10 mil establecimientos nacieron 61 y "
            "murieron 71 (INEGI, 2024b). Un control de inventario deficiente "
            "contribuye a este fenómeno al inmovilizar capital en mercancía sin "
            "rotación y al provocar pérdidas por caducidad o faltantes (INEGI, 2024b).")

    h2(doc, "2.2 Justificación e impacto potencial")
    p(doc, "La justificación de un sistema de gestión de inventarios para tiendas de "
            "abarrotes se sustenta en evidencia sobre la relación entre el control de "
            "existencias y el desempeño del negocio. La literatura especializada "
            "señala que una gestión adecuada de inventarios permite establecer "
            "políticas, métodos y procedimientos que reducen costos y mejoran la "
            "disponibilidad de productos, incidiendo directamente en la rentabilidad "
            "de las pymes (Romero-Agila et al., 2021). En el caso particular de las "
            "microempresas comerciales, el control de inventarios constituye una "
            "herramienta de información clave para la toma de decisiones gerenciales "
            "(Pin Figueroa et al., 2019).")
    p(doc, "El impacto potencial de implementar una herramienta sencilla y accesible "
            "es significativo. Al digitalizar el registro de productos, el comerciante "
            "puede detectar oportunamente los artículos con bajo stock, evitar la "
            "duplicación de registros y conocer el valor monetario inmovilizado en su "
            "inventario, lo que favorece decisiones de compra más informadas y una "
            "mejor planeación financiera (Romero-Agila et al., 2021). Asimismo, la "
            "adopción de tecnologías digitales accesibles se ha asociado con mayores "
            "probabilidades de supervivencia de las mipymes frente a entornos "
            "cambiantes (INEGI, 2024a).")

    h2(doc, "2.3 Identificación de stakeholders")
    p(doc, "El desarrollo de este sistema involucra a diversas partes interesadas "
            "(stakeholders) cuyo análisis permite dimensionar su alcance. A "
            "continuación se identifican los actores clave y la forma en que el "
            "proyecto los impacta (Romero-Agila et al., 2021).")
    tabla = doc.add_table(rows=1, cols=2)
    tabla.style = "Light Grid Accent 1"
    hdr = tabla.rows[0].cells
    hdr[0].paragraphs[0].add_run("Stakeholder").bold = True
    hdr[1].paragraphs[0].add_run("Interés / impacto").bold = True
    stakeholders = [
        ("Propietario(a) de la tienda",
         "Principal beneficiario: obtiene control del inventario, reduce pérdidas "
         "y toma mejores decisiones de compra."),
        ("Empleados / dependientes",
         "Usuarios operativos del sistema; agilizan el registro de ventas y "
         "reabastecimientos con menos errores."),
        ("Clientes de la tienda",
         "Beneficiarios indirectos: mayor disponibilidad de productos y menos "
         "desabastos."),
        ("Proveedores",
         "Se benefician de pedidos más precisos y oportunos basados en niveles "
         "reales de stock."),
        ("Desarrollador / estudiante",
         "Aplica y demuestra competencias técnicas de Python, POO y análisis de "
         "datos."),
    ]
    for nombre, interes in stakeholders:
        fila = tabla.add_row().cells
        fila[0].paragraphs[0].add_run(nombre).bold = True
        fila[1].text = interes
    doc.add_paragraph()

    h2(doc, "2.4 Objetivos")
    p(doc, "Objetivo general:", justificar=False)
    p(doc, "Desarrollar un sistema de gestión de inventarios en Python que permita a "
            "una tienda de abarrotes registrar, controlar y dar seguimiento a sus "
            "productos de forma sencilla, confiable y persistente, aplicando "
            "programación orientada a objetos y análisis de datos (Romero-Agila "
            "et al., 2021).")
    p(doc, "Objetivos específicos:", justificar=False)
    for obj in [
        "Implementar las operaciones básicas de gestión de productos (alta, "
        "reabastecimiento, venta y eliminación) mediante una arquitectura orientada "
        "a objetos que garantice la integridad de los datos.",
        "Incorporar mecanismos de control de datos que eviten productos duplicados "
        "y alerten sobre niveles de bajo stock, empleando estructuras de datos "
        "avanzadas como diccionarios y conjuntos.",
        "Asegurar la persistencia de la información mediante archivos JSON y "
        "demostrar buenas prácticas de análisis y limpieza de datos con la "
        "biblioteca pandas.",
    ]:
        par = doc.add_paragraph(style="List Number")
        par.add_run(obj)
        par.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    # ---------- 3. BASE DE DATOS, EDA Y LIMPIEZA ---------------------- #
    h1(doc, "3. Base de datos de ejemplo, análisis exploratorio y limpieza")

    h2(doc, "3.1 Descripción de la base de datos")
    p(doc, "Para ejemplificar el funcionamiento del sistema se construyó una base de "
            "datos de muestra (archivo inventario_muestra.csv) que simula el "
            "registro manual de 24 productos típicos de una tienda de abarrotes "
            "mexicana, con los campos nombre, precio, stock y categoría. De manera "
            "deliberada, la base contiene errores frecuentes de captura para "
            "demostrar la importancia del preprocesamiento de datos: espacios "
            "sobrantes, inconsistencias de mayúsculas y minúsculas, registros "
            "duplicados, valores faltantes, valores no numéricos y cantidades "
            "negativas (McKinney, 2010).")
    captura(doc, "contenido del archivo inventario_muestra.csv")

    h2(doc, "3.2 Análisis exploratorio de datos (EDA)")
    p(doc, "El análisis exploratorio constituye el primer paso para comprender la "
            "estructura, la calidad y las posibles anomalías de un conjunto de datos "
            "antes de utilizarlo. Mediante la biblioteca pandas se inspeccionaron las "
            "dimensiones del conjunto, los tipos de dato de cada columna, los valores "
            "faltantes, los registros duplicados y los estadísticos descriptivos de "
            "la variable precio (McKinney, 2010). El siguiente fragmento muestra la "
            "función de exploración implementada en el módulo analisis.py:")
    codigo(doc,
           "def explorar(df):\n"
           "    print(f\"Dimensiones: {df.shape}\")\n"
           "    print(df.dtypes)\n"
           "    print(df.isna().sum())           # valores faltantes\n"
           "    nombres = df['nombre'].str.strip().str.lower()\n"
           "    print(nombres.duplicated().sum())  # duplicados\n"
           "    print(pd.to_numeric(df['precio'], errors='coerce').describe())")
    p(doc, "La ejecución del análisis sobre la base de muestra arrojó los siguientes "
            "resultados, que evidencian la presencia de un nombre vacío, dos valores "
            "de stock faltantes y dos registros duplicados, además de una columna "
            "stock interpretada como texto debido a la presencia de un valor no "
            "numérico (McKinney, 2010):")
    codigo(doc,
           "Dimensiones (filas, columnas): (24, 4)\n"
           "Valores faltantes:  nombre 1 | precio 0 | stock 2 | categoria 0\n"
           "Filas duplicadas por nombre (normalizado): 2\n"
           "Estadisticos de 'precio': media 27.08 | min 11.00 | max 89.00")
    captura(doc, "salida completa del EDA en la consola")

    h2(doc, "3.3 Limpieza y preprocesamiento de datos")
    p(doc, "Tras el diagnóstico, se aplicó un pipeline de limpieza que transforma los "
            "datos crudos en un conjunto válido y consistente. El procedimiento "
            "normaliza los textos, convierte las columnas numéricas, imputa los "
            "valores de stock faltantes con cero, descarta los registros inválidos y "
            "elimina los duplicados conservando la primera aparición. Estas "
            "operaciones reflejan buenas prácticas de manejo de datos ampliamente "
            "recomendadas en la literatura de ciencia de datos (McKinney, 2010).")
    codigo(doc,
           "def limpiar(df):\n"
           "    df['nombre'] = df['nombre'].astype(str).str.strip()\n"
           "    df['precio'] = pd.to_numeric(df['precio'], errors='coerce')\n"
           "    df['stock']  = pd.to_numeric(df['stock'], errors='coerce').fillna(0)\n"
           "    df = df[df['precio'] > 0]              # descarta precios inválidos\n"
           "    df = df[df['stock'] >= 0]             # descarta stock negativo\n"
           "    df = df.drop_duplicates(subset='nombre_norm', keep='first')\n"
           "    return df")
    p(doc, "Como resultado del preprocesamiento, el conjunto se redujo de 24 a 20 "
            "registros válidos, descartando cuatro filas correspondientes a un "
            "registro sin nombre, un valor de stock negativo y dos duplicados. Este "
            "resultado confirma que cerca del 17 % de los registros capturados "
            "manualmente presentaban algún tipo de inconsistencia, una proporción que "
            "ilustra el riesgo de gestionar inventarios sin validaciones automáticas "
            "(Ramírez Castillejo & Hernández Pérez, 2023).")
    captura(doc, "tabla de datos limpios impresa en la consola")

    # --------------- 4. EVALUACIÓN DEL PROGRAMA ----------------------- #
    h1(doc, "4. Evaluación del programa: funciones y técnicas utilizadas")
    p(doc, "Esta sección describe las principales técnicas de programación empleadas, "
            "su función dentro del sistema y sus restricciones. El proyecto se "
            "estructuró en módulos independientes (modelo, inventario, análisis e "
            "interfaz) siguiendo el principio de separación de responsabilidades, lo "
            "que facilita el mantenimiento y la legibilidad del código (Van Rossum & "
            "Drake, 2009).")

    h2(doc, "4.1 Programación orientada a objetos (POO)")
    p(doc, "El sistema se modeló mediante dos clases principales. La clase Producto "
            "representa la unidad básica de información y encapsula sus atributos "
            "(nombre, precio, stock y categoría) junto con la validación de las "
            "reglas de negocio en su constructor. La clase Inventario actúa como "
            "contenedor y orquesta todas las operaciones sobre la colección de "
            "productos. El uso de POO permite reutilizar código y mantener la "
            "integridad del estado de los objetos (Van Rossum & Drake, 2009).")
    codigo(doc,
           "@dataclass\n"
           "class Producto:\n"
           "    nombre: str\n"
           "    precio: float\n"
           "    stock: int\n"
           "    categoria: str = 'General'\n"
           "    def __post_init__(self):\n"
           "        if self.precio <= 0:\n"
           "            raise ProductoInvalidoError('precio debe ser > 0')")
    p(doc, "Se empleó el decorador dataclass para reducir el código repetitivo de los "
            "métodos especiales y la propiedad property para exponer valores "
            "derivados, como el valor monetario del inventario. Como restricción, la "
            "validación se concentra en la creación del objeto; modificaciones "
            "posteriores directas a los atributos no se revalidan automáticamente, por "
            "lo que las operaciones se canalizan a través de métodos de la clase "
            "Inventario (Van Rossum & Drake, 2009).")

    h2(doc, "4.2 Estructuras de datos avanzadas")
    p(doc, "El control de productos se apoya en un diccionario (dict) que asocia una "
            "clave normalizada con cada objeto Producto, lo que permite acceder, "
            "actualizar o eliminar productos con una complejidad temporal constante, "
            "es decir O(1) en promedio. De forma complementaria, un conjunto (set) de "
            "claves se utiliza para verificar de manera inmediata la existencia de un "
            "producto y así impedir duplicados. Las listas se reservan para devolver "
            "resultados ordenados destinados a su visualización (Van Rossum & Drake, "
            "2009).")
    p(doc, "La principal restricción de estas estructuras es que residen en memoria "
            "volátil: su contenido se pierde al cerrar el programa. Por ello, el "
            "sistema las complementa con un mecanismo de persistencia en disco que se "
            "describe en la sección 5 (McKinney, 2010).")

    h2(doc, "4.3 Manipulación de datos con pandas y persistencia con JSON")
    p(doc, "La biblioteca pandas se utilizó para el análisis exploratorio y la "
            "limpieza de la base de datos de muestra, aprovechando su estructura "
            "DataFrame para operar sobre columnas completas de forma vectorizada y "
            "eficiente (McKinney, 2010). Para la persistencia se eligió el formato "
            "JSON por ser un estándar ligero, legible por humanos y nativo de Python "
            "a través del módulo json, lo que facilita guardar y recuperar el "
            "inventario entre sesiones sin necesidad de un sistema gestor de bases de "
            "datos (Van Rossum & Drake, 2009).")

    h2(doc, "4.4 Visualización con matplotlib")
    p(doc, "Para el análisis estadístico visual se empleó la biblioteca matplotlib, "
            "estándar de facto para la generación de gráficos en el ecosistema "
            "científico de Python (Hunter, 2007). Se generaron tres visualizaciones: "
            "un gráfico de barras de existencias por producto con resaltado del bajo "
            "stock, un gráfico de pastel de la distribución del valor por categoría y "
            "un histograma de la distribución de precios. Como restricción, las "
            "gráficas se generan con un backend no interactivo que las exporta como "
            "imágenes, decisión que favorece la automatización pero omite la "
            "interactividad en pantalla (Hunter, 2007).")

    # ---------- 5. COMPLEMENTACIÓN: CONTROL Y PERSISTENCIA ------------- #
    h1(doc, "5. Complementación: control de datos y persistencia")

    h2(doc, "5.1 Control de datos: duplicados y bajo stock")
    p(doc, "El control de datos garantiza la calidad de la información del inventario. "
            "Para evitar productos duplicados, antes de dar de alta un producto el "
            "sistema verifica si su clave normalizada ya existe en el conjunto de "
            "claves; en caso afirmativo, rechaza la operación y sugiere utilizar la "
            "función de reabastecimiento. Este mecanismo previene la fragmentación de "
            "la información, un problema común en registros manuales (Pin Figueroa "
            "et al., 2019).")
    codigo(doc,
           "def agregar_producto(self, producto):\n"
           "    if producto.clave in self._claves:      # control de duplicados\n"
           "        raise ProductoInvalidoError('el producto ya existe')\n"
           "    self._productos[producto.clave] = producto\n"
           "    self._claves.add(producto.clave)")
    p(doc, "Asimismo, el sistema identifica los productos cuyo nivel de existencias es "
            "igual o inferior a un umbral configurable (cinco unidades por defecto), "
            "generando una alerta de bajo stock. Esta funcionalidad permite al "
            "comerciante anticipar reabastecimientos y evitar desabastos que afectan "
            "las ventas (Pin Figueroa et al., 2019).")
    captura(doc, "reporte de productos con bajo stock (opción 6 del menú)")

    h2(doc, "5.2 Persistencia de datos")
    p(doc, "La persistencia asegura que el inventario permanezca disponible entre "
            "ejecuciones del programa. Cada operación que modifica el estado del "
            "inventario desencadena un guardado automático en el archivo "
            "inventario.json, y al iniciar el programa este se carga de forma "
            "automática. El proceso de carga incorpora manejo de excepciones para que "
            "un archivo inexistente o dañado no interrumpa la ejecución, iniciando en "
            "tal caso con un inventario vacío (Van Rossum & Drake, 2009).")
    codigo(doc,
           "def guardar(self):\n"
           "    datos = [p.a_dict() for p in self.listar()]\n"
           "    with open(self.ruta_archivo, 'w', encoding='utf-8') as f:\n"
           "        json.dump(datos, f, ensure_ascii=False, indent=4)")
    p(doc, "Esta estrategia de persistencia ligera resulta adecuada para el volumen "
            "de datos de una tienda pequeña y elimina la dependencia de "
            "infraestructura adicional, alineándose con la necesidad de soluciones "
            "accesibles para las microempresas (Romero-Agila et al., 2021).")

    h2(doc, "5.3 Análisis estadístico visual")
    p(doc, "Como complemento analítico, el sistema genera visualizaciones que "
            "resumen el estado del inventario y apoyan la toma de decisiones. La "
            "Figura 1 muestra las existencias por producto, resaltando en color los "
            "artículos por debajo del umbral de bajo stock (Hunter, 2007).")
    imagen(doc, "stock_por_producto.png", ancho=5.5)
    pie_figura(doc, "Figura 1. Existencias por producto con resaltado de bajo stock.")
    p(doc, "La Figura 2 presenta la distribución del valor monetario del inventario "
            "por categoría, información útil para identificar en qué tipo de productos "
            "se concentra el capital invertido del negocio (Hunter, 2007).")
    imagen(doc, "valor_por_categoria.png", ancho=4.5)
    pie_figura(doc, "Figura 2. Distribución del valor del inventario por categoría.")

    # ----------------- 6. INSTRUCCIONES DE USO ------------------------ #
    h1(doc, "6. Instrucciones de uso paso a paso")
    p(doc, "A continuación se describe el procedimiento para ejecutar e interactuar "
            "con el sistema. Se incluyen los comandos, una descripción del resultado "
            "esperado y espacios reservados para las capturas de pantalla "
            "correspondientes (Van Rossum & Drake, 2009).")

    pasos = [
        ("Paso 1. Instalar las dependencias",
         "Desde la carpeta del proyecto, ejecutar el siguiente comando para instalar "
         "las bibliotecas necesarias (pandas, matplotlib y python-docx):",
         "pip install -r requirements.txt",
         "instalación de dependencias completada"),
        ("Paso 2. Ejecutar el programa",
         "Iniciar la aplicación con el intérprete de Python. El sistema cargará "
         "automáticamente el inventario guardado y mostrará el menú principal:",
         "python src/main.py",
         "menú principal del sistema"),
        ("Paso 3. Mostrar el inventario",
         "Seleccionar la opción 1 para visualizar la tabla completa de productos, "
         "con su precio, stock, categoría y el valor total del inventario:",
         "Elige una opción: 1",
         "tabla del inventario"),
        ("Paso 4. Agregar un producto",
         "Seleccionar la opción 2 e ingresar los datos solicitados. Si el producto ya "
         "existe, el sistema lo notificará para evitar duplicados:",
         "Elige una opción: 2\nNombre del producto: Chocolate Abuelita\n"
         "Precio (MXN): 22.50\nStock inicial (unidades): 10\nCategoría: Bebidas",
         "alta de un producto nuevo"),
        ("Paso 5. Reabastecer y registrar ventas",
         "Las opciones 3 y 4 permiten sumar o descontar unidades del stock. El sistema "
         "valida que exista suficiente inventario antes de registrar una venta:",
         "Elige una opción: 3\nProducto a reabastecer: Frijol Bayo Kg\n"
         "Unidades a agregar: 20",
         "reabastecimiento de un producto"),
        ("Paso 6. Consultar bajo stock",
         "La opción 6 lista los productos cuyo nivel de existencias requiere atención, "
         "facilitando la planeación de compras:",
         "Elige una opción: 6",
         "reporte de bajo stock"),
        ("Paso 7. Cargar la base de muestra y generar gráficas",
         "Las opciones 7 y 8 permiten cargar la base de datos de ejemplo (ejecutando "
         "el EDA y la limpieza) y generar las gráficas estadísticas en la carpeta "
         "graficas/:",
         "Elige una opción: 7\n...\nElige una opción: 8",
         "gráficas generadas"),
    ]
    for titulo, descripcion, cmd, cap in pasos:
        h2(doc, titulo)
        p(doc, descripcion)
        codigo(doc, cmd)
        captura(doc, cap)

    # ----------------------- 7. CONCLUSIÓN ---------------------------- #
    h1(doc, "7. Conclusión")
    p(doc, "El desarrollo de este sistema de gestión de inventarios permitió integrar "
            "y aplicar de manera práctica los temas revisados durante la "
            "certificación en Python: desde los fundamentos del lenguaje y las "
            "estructuras de control, hasta la programación orientada a objetos, las "
            "estructuras de datos avanzadas, la manipulación de datos con pandas y la "
            "visualización con matplotlib. La construcción modular del proyecto "
            "evidenció la importancia de las buenas prácticas de codificación, como la "
            "separación de responsabilidades, la documentación mediante docstrings y "
            "el manejo controlado de excepciones (Van Rossum & Drake, 2009).")
    p(doc, "Más allá del ejercicio técnico, el proyecto reafirma el valor de la "
            "programación como herramienta para resolver problemas reales. Python, "
            "por su accesibilidad y su robusto ecosistema de bibliotecas, demuestra "
            "ser un lenguaje idóneo para construir soluciones que, sin requerir "
            "infraestructura costosa, aportan valor tangible a contextos cotidianos "
            "(McKinney, 2010). La capacidad de transformar datos crudos y "
            "desordenados en información confiable es, precisamente, una de las "
            "competencias más demandadas en la actualidad (Hunter, 2007).")
    p(doc, "En el ámbito de las microempresas, sistematizar la gestión de inventarios "
            "trasciende la mera digitalización: incide en la rentabilidad, reduce "
            "pérdidas y fortalece la toma de decisiones, factores asociados con una "
            "mayor probabilidad de supervivencia empresarial en un entorno donde la "
            "mortalidad de los negocios es elevada (INEGI, 2024b). Por ello, "
            "soluciones accesibles como la aquí presentada tienen el potencial de "
            "contribuir, aunque sea de forma modesta, a la profesionalización y "
            "permanencia de las tiendas de abarrotes mexicanas (Romero-Agila et al., "
            "2021).")

    # ----------------------- 8. REFERENCIAS --------------------------- #
    h1(doc, "8. Referencias")
    referencias = [
        "Hunter, J. D. (2007). Matplotlib: A 2D graphics environment. Computing in "
        "Science & Engineering, 9(3), 90–95. https://doi.org/10.1109/MCSE.2007.55",

        "Instituto Nacional de Estadística y Geografía. (2024a). Estadísticas a "
        "propósito del Día de las Micro, Pequeñas y Medianas Empresas (Comunicado de "
        "prensa núm. 383/24). INEGI. "
        "https://www.inegi.org.mx/contenidos/saladeprensa/aproposito/2024/EAP_MIPYMES24.pdf",

        "Instituto Nacional de Estadística y Geografía. (2024b). Estudio sobre la "
        "Demografía de los Negocios (EDN) 2023 (Comunicado de prensa núm. 68/24). "
        "INEGI. "
        "https://www.inegi.org.mx/contenidos/saladeprensa/boletines/2024/EDN/EDN2023.pdf",

        "McKinney, W. (2010). Data structures for statistical computing in Python. "
        "En Proceedings of the 9th Python in Science Conference (pp. 56–61). "
        "https://doi.org/10.25080/Majora-92bf1922-00a",

        "Pin Figueroa, F. E., Pin Figueroa, M. A., & Pin Figueroa, J. L. (2019). El "
        "control de los inventarios y su incidencia en las decisiones gerenciales en "
        "las microempresas de comercio de Jipijapa. FIPCAEC, 4(1 Especial), 1–20. "
        "https://www.fipcaec.com/index.php/fipcaec/article/view/106",

        "Ramírez Castillejo, M. de los Á., & Hernández Pérez, J. (2023). La ausencia "
        "de la gestión administrativa en las pymes de México, una realidad que les "
        "impide crecer. Revista IPSUMTEC, 6(2). "
        "https://revistas.milpaalta.tecnm.mx/index.php/IPSUMTEC/article/view/137",

        "Romero-Agila, M. L., Erazo-Álvarez, J. C., Narváez-Zurita, C. I., & "
        "Matovelle-Romo, M. (2021). La gestión de inventarios en las PYMES del sector "
        "de la construcción. Polo del Conocimiento, 6(9), 100–124. "
        "https://doi.org/10.23857/pc.v6i9.3124",

        "Van Rossum, G., & Drake, F. L. (2009). The Python language reference manual. "
        "Python Software Foundation. https://docs.python.org/3/reference/",
    ]
    for ref in referencias:
        par = doc.add_paragraph()
        par.paragraph_format.left_indent = Inches(0.5)
        par.paragraph_format.first_line_indent = Inches(-0.5)  # sangría francesa
        par.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        par.add_run(ref).font.size = Pt(10)

    nota = doc.add_paragraph()
    nota.add_run(
        "Nota sobre las fuentes: las cifras oficiales provienen de comunicados del "
        "INEGI verificados directamente en los documentos primarios. Las referencias "
        "técnicas (pandas, matplotlib, Python) corresponden a las publicaciones "
        "originales de cada herramienta. Las cifras de demografía empresarial "
        "corresponden al periodo mayo 2019–mayo 2023 reportado por el EDN 2023."
    ).italic = True
    nota.runs[0].font.size = Pt(9)
    nota.runs[0].font.color.rgb = GRIS

    os.makedirs(os.path.dirname(SALIDA), exist_ok=True)
    doc.save(SALIDA)
    print(f"Informe generado: {SALIDA}")


if __name__ == "__main__":
    construir()
